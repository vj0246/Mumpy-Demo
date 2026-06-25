"""
docstore.py
-----------
Lets a user upload a PDF / Word / text document (e.g. an annual report or a
financial statement) and then ask questions about it in the chat.

Kept deliberately simple and dependency-light so it stays on the demo's "no new
paid keys" rule:

  * text is extracted locally (pypdf for PDF, python-docx for .docx)
  * the text is split into overlapping chunks
  * retrieval is plain keyword scoring (no embeddings / vector DB)

The LLM call that turns retrieved chunks into an answer lives in chat_agent.py,
so it reuses the same Groq model. Documents are held in memory, keyed by the
chat's thread id (one active document per chat — a new upload replaces it).
"""

import io
import re

# thread_id -> {"name": str, "text": str, "chunks": [str, ...], "chars": int}
_DOCS: dict = {}

_MAX_CHARS = 400_000          # guard against an enormous paste/upload
_CHUNK = 1100                 # chunk size (chars); financial tables stay mostly intact
_OVERLAP = 200


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def extract_text(filename: str, raw: bytes) -> str:
    """Pull plain text out of a PDF, Word (.docx) or text file."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except Exception as e:  # pragma: no cover - dependency guard
            raise RuntimeError("PDF support needs 'pypdf' (pip install pypdf).") from e
        reader = PdfReader(io.BytesIO(raw))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(parts)

    elif name.endswith(".docx"):
        try:
            import docx  # python-docx
        except Exception as e:  # pragma: no cover - dependency guard
            raise RuntimeError("Word support needs 'python-docx' (pip install python-docx).") from e
        d = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:                       # financial statements live in tables
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)

    elif name.endswith((".txt", ".md", ".csv")):
        text = raw.decode("utf-8", errors="ignore")

    elif name.endswith(".doc"):
        raise RuntimeError("Legacy .doc isn't supported — save as .docx or PDF and re-upload.")

    else:
        raise RuntimeError("Unsupported file type. Upload a PDF, Word (.docx) or text file.")

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise RuntimeError("Couldn't read any text from that file (it may be scanned/image-only).")
    return text[:_MAX_CHARS]


def _chunk(text: str) -> list:
    chunks, i, n = [], 0, len(text)
    while i < n:
        chunks.append(text[i:i + _CHUNK])
        i += _CHUNK - _OVERLAP
    return chunks


def add_document(thread_id: str, filename: str, text: str) -> dict:
    chunks = _chunk(text)
    _DOCS[thread_id] = {"name": filename, "text": text, "chunks": chunks, "chars": len(text)}
    return {"name": filename, "chars": len(text), "chunks": len(chunks)}


def clear(thread_id: str) -> None:
    _DOCS.pop(thread_id, None)


def has_document(thread_id: str) -> bool:
    return thread_id in _DOCS


def doc_name(thread_id: str):
    d = _DOCS.get(thread_id)
    return d["name"] if d else None


_WORD = re.compile(r"[a-z0-9]+")


def _tokens(s: str) -> list:
    return _WORD.findall(s.lower())


def retrieve(thread_id: str, query: str, k: int = 5) -> str:
    """Return the most relevant chunks for the query as one context string.

    Keyword overlap scoring with a light bonus for chunks that contain digits
    (financial questions usually want the numbers). Falls back to the opening
    chunks if nothing matches.
    """
    doc = _DOCS.get(thread_id)
    if not doc:
        return ""
    q_terms = set(_tokens(query))
    chunks = doc["chunks"]
    if not q_terms:
        return "\n\n---\n\n".join(chunks[:k])

    scored = []
    for idx, ch in enumerate(chunks):
        toks = _tokens(ch)
        if not toks:
            continue
        tset = set(toks)
        overlap = sum(1 for t in q_terms if t in tset)
        if overlap == 0:
            continue
        digit_bonus = 0.5 if any(c.isdigit() for c in ch) else 0.0
        scored.append((overlap + digit_bonus, idx, ch))

    if not scored:
        return "\n\n---\n\n".join(chunks[:k])
    scored.sort(key=lambda x: (-x[0], x[1]))
    top = sorted(scored[:k], key=lambda x: x[1])      # restore document order
    return "\n\n---\n\n".join(ch for _, _, ch in top)
